"""Leitura do CV (PDF, DOCX, MD, TXT) e extracao de fatos.

Principio: o parser EXTRAI, nunca INFERE. Se uma secao nao for encontrada,
ela fica vazia — jamais preenchida com suposicao.
"""
from __future__ import annotations

import re
from pathlib import Path

from app.crawler.extract import detect_technologies, normalize_text
from app.logging_setup import get_logger
from app.models.profile import ResumeFacts

log = get_logger("resume.parser")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".markdown"}

#: Cabecalhos de secao em PT e EN.
SECTION_HEADINGS: dict[str, list[str]] = {
    "experience": ["experiencia profissional", "experiencia", "professional experience",
                   "work experience", "employment", "historico profissional", "carreira"],
    "education": ["educacao", "formacao academica", "formacao", "education", "academic"],
    "certifications": ["certificacoes", "certificados", "certifications", "licenses",
                       "cursos", "courses"],
    "languages": ["idiomas", "languages", "linguas"],
    "projects": ["projetos", "projects", "portfolio", "side projects"],
    "achievements": ["conquistas", "achievements", "realizacoes", "premios", "awards",
                     "destaques", "highlights"],
    "skills": ["habilidades", "competencias", "skills", "tecnologias", "technologies",
               "technical skills", "stack", "conhecimentos"],
}


def extract_text(path: Path) -> str:
    """Texto cru do arquivo. Erro claro se o formato nao for suportado."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("Para ler PDF instale as dependencias: pip install pypdf") from exc
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        try:
            import docx
        except ImportError as exc:
            raise RuntimeError("Para ler DOCX instale: pip install python-docx") from exc
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    if suffix in {".md", ".markdown", ".txt"}:
        return path.read_text(encoding="utf-8", errors="replace")
    raise RuntimeError(
        f"Formato '{suffix}' nao suportado. Use PDF, DOCX, MD ou TXT."
    )


def _is_heading(line: str) -> str | None:
    """A linha e um cabecalho de secao conhecido? Devolve a chave da secao."""
    stripped = line.strip().strip("#*_-:•").strip()
    if not stripped or len(stripped) > 60:
        return None
    normalized = normalize_text(stripped)
    for key, variants in SECTION_HEADINGS.items():
        for variant in variants:
            if normalized == variant or normalized.rstrip(":") == variant:
                return key
    return None


def split_sections(text: str) -> dict[str, list[str]]:
    """Divide o CV em secoes pelos cabecalhos reconhecidos."""
    sections: dict[str, list[str]] = {key: [] for key in SECTION_HEADINGS}
    sections["_preamble"] = []
    current = "_preamble"
    for line in text.splitlines():
        heading = _is_heading(line)
        if heading:
            current = heading
            continue
        content = line.strip().strip("•*-–— ").strip()
        if content:
            sections[current].append(content)
    return sections


def _clean_entries(lines: list[str], min_len: int = 3, max_items: int = 40) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for line in lines:
        collapsed = re.sub(r"\s+", " ", line).strip()
        if len(collapsed) < min_len:
            continue
        key = normalize_text(collapsed)
        if key in seen:
            continue
        seen.add(key)
        out.append(collapsed)
    return out[:max_items]


def parse_resume(path: str | Path) -> ResumeFacts:
    """Extrai fatos do CV. Somente o que esta escrito no arquivo."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Curriculo nao encontrado: {p}")
    if p.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise RuntimeError(
            f"Extensao '{p.suffix}' nao suportada. Use: "
            + ", ".join(sorted(SUPPORTED_EXTENSIONS))
        )

    text = extract_text(p)
    sections = split_sections(text)

    # Tecnologias: detectadas no documento inteiro, com vocabulario canonico.
    technologies = detect_technologies(text)

    facts = ResumeFacts(
        source_file=str(p),
        raw_text=text,
        technologies=technologies,
        experience_entries=_clean_entries(sections["experience"], min_len=8),
        education=_clean_entries(sections["education"], min_len=5),
        certifications=_clean_entries(sections["certifications"], min_len=3),
        languages=_clean_entries(sections["languages"], min_len=3, max_items=12),
        projects=_clean_entries(sections["projects"], min_len=8),
        achievements=_clean_entries(sections["achievements"], min_len=8),
    )
    log.info("CV lido: %s (%d chars, %d tecnologias detectadas)",
             p.name, len(text), len(technologies))
    return facts


def find_resume(resume_dir: str | Path) -> Path | None:
    """Escolhe o CV mais recente da pasta de curriculos."""
    directory = Path(resume_dir)
    if not directory.exists():
        return None
    candidates = [
        f for f in directory.iterdir()
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
        and not f.name.startswith(".")
    ]
    if not candidates:
        return None
    return max(candidates, key=lambda f: f.stat().st_mtime)
